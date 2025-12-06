from Models.DocumentModel import DocumentModel
from database import AsyncSessionLocal
from schemas.DocumentSchema import DocumentCreate, DocumentRead, StatusEnum
from Exception.IngestionException import SaveDocumentError,UpdateDocumentStatusError,UnsupportedFileTypeError,DocumentFailedError
from Exception.DocumentException import DocumentReadError,NumberPageError,DocumentDeleteError
from sqlalchemy.exc import SQLAlchemyError
from docx import Document as DocxDocument
from sqlalchemy import select
from docx2pdf import convert
from pathlib import Path
from fastapi import UploadFile
import unicodedata
import re
import uuid
import os
import PyPDF2
import tempfile
from typing import List
import aiofiles
import platform
from pathlib import Path
import zipfile

BASE_DIR = Path(__file__).resolve().parent.parent

#-------crud------------

async def delete_all_document(chat_id:uuid.UUID):
    """
    Delete a chat's document from the database.

    Args:
        doc (DocumentRead): The document object to delete. 
                            Must contain at least the `id` attribute.

    Raises:
        DocumentDeleteError: If a SQLAlchemy error occurs during deletion.

    Notes:
        - Uses an asynchronous session with `AsyncSessionLocal`.
        - Deletion is performed within a transactional context (`session.begin()`), 
          which ensures automatic commit.
        - If the document does not exist, no action is taken.
    """
    try:
        async with AsyncSessionLocal() as session:
            async with session.begin():

                stmt = select(DocumentModel).where(DocumentModel.chat_id == chat_id)
                response = await session.execute(stmt)
                object = response.scalars().all()

                if object :
                    for doc in object:
                        await session.delete(doc)

    except SQLAlchemyError as e:
        raise DocumentDeleteError("Failed to delete docs")
    
    
async def delete_document(doc: DocumentRead):
    """
    Delete a document from the database based on its ID.

    Args:
        doc (DocumentRead): The document object to delete. 
                            Must contain at least the `id` attribute.

    Raises:
        DocumentDeleteError: If a SQLAlchemy error occurs during deletion.

    Notes:
        - Uses an asynchronous session with `AsyncSessionLocal`.
        - Deletion is performed within a transactional context (`session.begin()`), 
          which ensures automatic commit.
        - If the document does not exist, no action is taken.
    """
    try:
        async with AsyncSessionLocal() as session:
            async with session.begin():
                
                stmt = select(DocumentModel).where(DocumentModel.id == doc.id)
                response = await session.execute(stmt)
                object = response.scalars().first()

                if object:
                    await session.delete(object)
                    
    except SQLAlchemyError:
        raise DocumentDeleteError("Failed to delete the docs")
    
    
async def delete_documents_except(chat_id: uuid.UUID, keep_ids: List[uuid.UUID]) -> int:
    """
    Delete all documents for a given chat except those whose IDs are in the keep_ids list.

    This function opens an async SQLAlchemy session internally, filters documents by chat_id,
    and deletes all documents whose IDs are not in the keep_ids list.

    Args:
        chat_id (uuid.UUID): The chat ID to filter documents.
        keep_ids (List[uuid.UUID]): List of document IDs to keep.

    Returns:
        int: Number of documents deleted.

    Raises:
        HTTPException 500: If deletion fails due to a database error.
    """
    try:
        async with AsyncSessionLocal() as session:
            async with session.begin():
                stmt = select(DocumentModel).where(
                    DocumentModel.chat_id == chat_id,
                    DocumentModel.id.not_in(keep_ids)
                )
                result = await session.execute(stmt)
                objects = result.scalars().all()

                for obj in objects:
                    await session.delete(obj)

    except SQLAlchemyError as e:
        raise e



async def create_document(document: DocumentCreate) -> DocumentRead :
    """
    Creates and saves a document in the database.

    Args:
        document (DocumentCreate): The document data to save.

    Returns:
        DocumentRead: The saved document as a read schema.

    Raises:
        SaveDocumentError: If saving the document fails.
    """
    try :
        async with AsyncSessionLocal() as session : #start communication with the db
            async with session.begin() :    
                doc_title = Path(document.meta_data.get('file_name', '')).stem 
                new_document = DocumentModel(
                    chat_id=document.chat_id,
                    title=doc_title,
                    text=document.text,
                    meta_data=document.meta_data
                )
                session.add(new_document)
                await session.flush() #To give the object an id

                return DocumentRead.from_orm(new_document)
            
    except (SQLAlchemyError,TypeError,SyntaxError) as e :
        raise SaveDocumentError(f"Failed to create document {document.meta_data.get('file_name')}.OriginalError : {e}") from e


async def update_document_status(document_id: uuid.UUID, new_status: StatusEnum) :
    """
    Updates the status of a document.

    Args:
        document_id (uuid.UUID): The document's ID.
        new_status (StatusEnum): The new status to set.

    Raises:
        UpdateDocumentStatusError: If updating the status fails.
    """
    try :
        async with AsyncSessionLocal() as session: #start communication with the db
            async with session.begin() : #start the conversation
                stmt = select(DocumentModel).where(DocumentModel.id == document_id)
                result = await session.execute(stmt)
                document_obj = result.scalars().first()#tansofrm the first line of data as an python object
                if not document_obj :
                    raise Exception(f"Document id {document_id} not found")
                document_obj.status = new_status #change the title value
    except (SQLAlchemyError,TypeError,SyntaxError) as e :
        raise UpdateDocumentStatusError(f'Failed to update status of  document {document_id}.Original Error {e}') from e


async def read_document(chat_id : uuid.UUID ) -> List[DocumentRead] :
    """
    Retrieves all documents associated with a chat.

    Args:
        chat_id (uuid.UUID): The chat ID.

    Returns:
        List[DocumentRead]: A list of documents linked to the chat.

    Raises:
        DocumentReadError: If reading documents fails.
    """
    try :
        async with AsyncSessionLocal() as session : #start communication with the db
            async with session.begin() : #start the conversation
                stmt = select(DocumentModel).where(DocumentModel.chat_id == chat_id)
                result = await session.execute(stmt)
                documents = result.scalars().all() #return all the datas fom the db as an python object
                return [DocumentRead.from_orm(doc) for doc in documents]
    except Exception as e :
        raise DocumentReadError(f"Error while Reading documents.Cause : {e}")


async def read_document_id(id : uuid.UUID) -> DocumentRead :
    """
    Retrieves a document by its ID.

    Args:
        id (uuid.UUID): The document ID.

    Returns:
        DocumentRead: The document.

    Raises:
        DocumentReadError: If reading the document fails.
    """
    try:
        async with AsyncSessionLocal() as session : #start communication with the db
            async with session.begin() : #start the conversation 
                stmt = select(DocumentModel).where(DocumentModel.id == id)
                result = await session.execute(stmt)
                document = result.scalars().first()
                return DocumentRead.from_orm(document)
    except Exception as e :
        raise DocumentReadError(f"Error while Reading documents.Cause : {e}")


#--------- other functions ------

def get_num_pages_pdf(file_path: str) -> int :
    """
    Returns the number of pages in a PDF file.

    Args:
        file_path (str): Path to the PDF file.

    Returns:
        int: Number of pages in the PDF.

    Raises:
        NumberPageError: If the PDF has more than 150 pages.
    """
    with open(file_path, "rb") as f : #interacte with a file
        reader = PyPDF2.PdfReader(f)
        num_pages = len(reader.pages)
        if num_pages > 150 :
            raise NumberPageError(f"{file_path} has more than 150 pages ({num_pages})")
        return num_pages


def get_num_pages_docx(file_path: str) -> int:
    """
    Returns the number of pages in a DOCX file.

    On Windows, it converts the DOCX to a temporary PDF and counts PDF pages.
    On other systems, it estimates pages based on paragraph count (~40 per page).

    Args:
        file_path (str): Path to the DOCX file.

    Returns:
        int: Number of pages in the DOCX file.

    Raises:
        NumberPageError: If the page count cannot be determined or exceeds the limit.
    """
    if platform.system() == "Windows":
        try:
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_pdf:
                convert(file_path, tmp_pdf.name)  # transform the docx into pdf
                num_pages = get_num_pages_pdf(tmp_pdf.name)
            return num_pages
        finally:
            if os.path.exists(tmp_pdf.name):
                os.remove(tmp_pdf.name)  # delete the temporary .pdf file
    else:
        try:
            doc = DocxDocument(file_path)
            total_paragraphs = len(doc.paragraphs)
            estimated_pages = max(1, total_paragraphs // 40)

            if estimated_pages > 150:
                raise NumberPageError(
                    f"{file_path} has more than 150 pages (estimated: {estimated_pages})"
                )

            return estimated_pages
        except Exception as e:
            raise NumberPageError(f"Could not estimate pages for {file_path}: {e}")


def extract_text_pdf(file_path: Path) -> str :
    """
    Extracts text content from a PDF file.

    Args:
        file_path (Path): Path to the PDF file.

    Returns:
        str: Extracted text content.
    """
    text = ""
    with open(file_path, "rb") as f :#Interacte with the file
        reader = PyPDF2.PdfReader(f) #get the f file as a PdfReader object
        for page in reader.pages :
            text += page.extract_text() or ""
    return text


def extract_text_pdf_with_pages(file_path: Path) -> tuple[str, dict[int, tuple[int, int]]]:
    """
    Extracts text content from a PDF file with page mapping.
    
    Args:
        file_path (Path): Path to the PDF file.
    
    Returns:
        tuple: (full_text, page_map)
            - full_text (str): Complete extracted text
            - page_map (dict): Maps character positions to page numbers
                {page_num: (start_char_pos, end_char_pos)}
    """
    text = ""
    page_map = {}
    
    with open(file_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page_num, page in enumerate(reader.pages, start=1):
            start_pos = len(text)
            page_text = page.extract_text() or ""
            text += page_text
            end_pos = len(text)
            page_map[page_num] = (start_pos, end_pos)
    
    return text, page_map


def extract_text_docx(file_path: Path) -> str :
    """
    Extracts text content from a DOCX file.

    Args:
        file_path (Path): Path to the DOCX file.

    Returns:
        str: Extracted text content.
        
    Raises:
        ValueError: If the DOCX file is corrupted or invalid.
    """
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    if not zipfile.is_zipfile(file_path):
        raise ValueError(f"Invalid DOCX file - not a valid ZIP archive: {file_path.name}")
    
    try:
        doc = DocxDocument(file_path)
        return "\n".join([p.text for p in doc.paragraphs])
        
    except KeyError as e:
        raise ValueError(f"Corrupted DOCX file - missing component {e} in: {file_path.name}")
        
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX {file_path.name}: {type(e).__name__}: {e}")


def extract_text_docx_with_pages(file_path: Path) -> tuple[str, dict[int, tuple[int, int]]]:
    """
    Extracts text content from a DOCX file with estimated page mapping.
    
    Args:
        file_path (Path): Path to the DOCX file.
    
    Returns:
        tuple: (full_text, page_map)
            - full_text (str): Complete extracted text
            - page_map (dict): Maps character positions to estimated page numbers
    
    Note:
        Page estimation is approximate based on paragraph count (~40 paragraphs per page).
    """
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    if not zipfile.is_zipfile(file_path):
        raise ValueError(f"Invalid DOCX file - not a valid ZIP archive: {file_path.name}")
    
    try:
        doc = DocxDocument(file_path)
        text = ""
        page_map = {}
        
        current_page = 1
        paragraphs_in_page = 0
        page_start = 0
        
        for paragraph in doc.paragraphs:
            para_text = paragraph.text + "\n"
            text += para_text
            paragraphs_in_page += 1
            
            # Estimate: ~40 paragraphs per page
            if paragraphs_in_page >= 40:
                page_map[current_page] = (page_start, len(text))
                current_page += 1
                page_start = len(text)
                paragraphs_in_page = 0
        
        # Add last page
        if paragraphs_in_page > 0:
            page_map[current_page] = (page_start, len(text))
        
        return text, page_map
        
    except KeyError as e:
        raise ValueError(f"Corrupted DOCX file - missing component {e} in: {file_path.name}")
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX {file_path.name}: {type(e).__name__}: {e}")


def load_file_from_uploads(target_folder: Path = Path("App/Data/Uploads")) -> List[dict] :
    """
    Loads files from the uploads directory and extracts text content with page mapping.

    Args:
        target_folder (Path): The folder to load files from.

    Returns:
        List[dict]: List of dictionaries representing each file with text, metadata, and page_map.

    Raises:
        UnsupportedFileTypeError: If a file type is not supported.
    """
    Documents = [] # dict list
    try :
        for file_path in target_folder.iterdir() :
            if file_path.is_file() :
                ext = file_path.suffix.lower()
                if ext == ".pdf"  :
                    text, page_map = extract_text_pdf_with_pages(file_path)
                    text = clean_text(text)
                elif ext == ".docx" :
                    text, page_map = extract_text_docx_with_pages(file_path)
                    text = clean_text(text)
                else : #the doc file is neither /pdf or docx file
                    raise UnsupportedFileTypeError(f"Document type not supported: {ext}")
                doc_dict = {
                    "text": text,
                    "file_extension": ext,
                    "file_name": file_path.name,
                    "file_path": file_path,
                    "page_map": page_map  # NOUVEAU: ajout du mapping des pages
                }
                Documents.append(doc_dict) # add the doc_dict to the list
        return Documents
    except Exception as e :
        raise e


def delete_file_from_uploads() :
    """
    Deletes all files from the uploads directory.
    """
    target_folder = "App/Data/Uploads"
    for filename in os.listdir(target_folder) :
        file_path = os.path.join(target_folder, filename) #create a way to the files
        if os.path.isfile(file_path) :
            os.remove(file_path) #delete the file


def clean_text(text: str) -> str:
    """
    Cleans text by removing unnecessary characters and normalizing Unicode.

    Args:
        text (str): Input text.

    Returns:
        str: Cleaned text.
    """
    text = unicodedata.normalize("NFKC", text)  # Normalize Unicode characters to NFKC form
    text = re.sub(r"[\x00-\x1F\x7F-\x9F\u2000-\u206F\u2E00-\u2E7F\uF000-\uFFFF]", " ", text)  # Remove control characters and miscellaneous Unicode punctuation
    text = text.replace("\n", " ").replace("\r", " ")  # Replace newline and carriage return with spaces
    text = re.sub(r"\s+", " ", text) # Collapse multiple spaces into one
    text = text.strip()  # Remove leading and trailing whitespace
    return text

async def charge_document_uploads_directory(files: List[UploadFile]):
    """
    Create Files and charge them into the uploads directory.

    Args:
        files(list[UploadFile])
    """
    try:
        target_folder = "App/Data/Uploads"
        os.makedirs(target_folder, exist_ok=True)  # if the corresponding dir exists do not create it

        for file in files:
            file_path = os.path.join(target_folder, file.filename)  # create a file_path corresponding to each file

            # Ensure the file pointer is at the beginning
            await file.seek(0)
            
            async with aiofiles.open(file_path, "wb") as f:  # create a way to write into the files
                content = await file.read()
                await f.write(content)

    except Exception as e:
        raise e


async def if_exist(chat_id : uuid.UUID) -> bool :
    """
    Checks if any document is linked to a chat.

    Args:
        chat_id (uuid.UUID): The chat ID to check for linked documents.

    Returns:
        bool: True if at least one document is linked, False otherwise.
    """
    documents = await read_document(chat_id)

    if not documents :
        return False
    else : 
        return True


async def add_metadata_to_docs(documents, chat_id: uuid.UUID) -> List[DocumentCreate] :
    """
    Adds metadata to each document and converts it into DocumentCreate objects.
    Now includes page_map in metadata.

    Args:
        documents (list): List of document dictionaries with keys 'text', 'file_name', 'file_extension', 'file_path', 'page_map'.
        chat_id (uuid.UUID): The chat ID to associate with each document.

    Returns:
        List[DocumentCreate]: List of DocumentCreate objects with metadata added.

    Raises:
        NumberPageError: If a document has more than 150 pages.
        UnsupportedFileTypeError: If a document has an unsupported file type.
    """
    docs: List[DocumentCreate] = []
    
    for doc in documents :
        try :
            text = doc['text']
            file_name = doc["file_name"]
            file_ext = doc["file_extension"]
            file_path = doc["file_path"]
            page_map = doc.get("page_map", {})  # NOUVEAU: récupérer le page_map
            
            if file_ext == ".pdf":
                num_pages = get_num_pages_pdf(file_path) #get the number of page of each file
            elif file_ext == ".docx":
                num_pages  = get_num_pages_docx(file_path) #get the number of page of each file
            else:
                raise UnsupportedFileTypeError(f"Document type not supported: {file_ext}")
            

            meta_data ={
                "file_name": file_name,
                "file_ext":file_ext,
                "num_pages":num_pages,
                "page_map": page_map  # NOUVEAU: ajout du mapping des pages
            }
            new_doc = DocumentCreate(
                chat_id= chat_id,
                text= text,
                meta_data= meta_data
            )
             #convert the dict into DocumentCreate object
            docs.append(new_doc)

        except NumberPageError as e :
            raise NumberPageError(f"Doc : {file_name} has more than 150 pages") 
        
        except Exception as e:
            raise 

    return docs


async def set_chunked_document(document_id: uuid.UUID) :
    """
    Sets the status of a document to 'chunked'.

    Args:
        document_id (uuid.UUID): The document ID to update.

    Raises:
        UpdateDocumentStatusError: If updating the document status fails.
    """
    try:
        await update_document_status(document_id, StatusEnum.chunked)
    except UpdateDocumentStatusError as e:
        raise


async def set_ready_document(document_id: uuid.UUID) :
    """
    Sets the status of a document to 'ready'.

    Args:
        document_id (uuid.UUID): The document ID to update.

    Raises:
        UpdateDocumentStatusError: If updating the document status fails.
    """
    try :
        await update_document_status(document_id, StatusEnum.ready)
    except UpdateDocumentStatusError as e :
        raise 


async def set_document_failed(document_id: uuid.UUID) ->DocumentRead :
    """
    Sets the status of a document to 'failed'.

    Args:
        document_id (uuid.UUID): The document ID to update.

    Raises:
        UpdateDocumentStatusError: If updating the document status fails.
    """
    try :
        await update_document_status(document_id, StatusEnum.failed)
    except UpdateDocumentStatusError as e :
        raise       


async def compare_name(file_name: str, chat_id: uuid.UUID) -> DocumentRead | None:
    """
    compare the database's file with anothe other files 

    Args:
        file_name(str): the file that is being uploading
        chat_id(uuid.UUID): the chat's ID
    
    Returns:
        None | DocumentRead : if it is already uploaded or not

    Raises:
        Exception: If it fails
    """
    try:
        docs: list[DocumentRead] = await read_document(chat_id)
        for doc in docs:
            if file_name == doc.meta_data["file_name"]:
                return doc
        return None
    
    except Exception as e:
        raise e